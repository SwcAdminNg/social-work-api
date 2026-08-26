import io
import json
import re
from typing import Any

import httpx
from docx import Document
from fastapi import HTTPException, status
from pydantic import BaseModel, Field, ValidationError
from pypdf import PdfReader

from app.core.config import settings
from app.modules.course.content_dto import AssessmentAIProviderEnum, QuizOptionCreateDTO, QuizQuestionCreateDTO
from app.modules.course.content_entity import MultiAnswerModeEnum


PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class RawQuizAIGenerationResult(BaseModel):
    questions: list[QuizQuestionCreateDTO] = Field(default_factory=list)


class QuizAIGenerationResult(RawQuizAIGenerationResult):
    provider: AssessmentAIProviderEnum
    model: str


def extract_assessment_text(file_name: str, content_type: str | None, data: bytes) -> str:
    if not data:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "Uploaded assessment document is empty")

    if len(data) > settings.assessment_ai_max_file_size_bytes:
        max_mb = settings.assessment_ai_max_file_size_bytes // (1024 * 1024)
        raise HTTPException(status.HTTP_413_REQUEST_ENTITY_TOO_LARGE, f"Assessment document must be {max_mb}MB or smaller")

    extension = _file_extension(file_name)
    normalized_content_type = (content_type or "").split(";")[0].strip().lower()
    is_pdf = extension == ".pdf" or normalized_content_type == PDF_CONTENT_TYPE
    is_docx = extension == ".docx" or normalized_content_type == DOCX_CONTENT_TYPE

    if is_pdf:
        text = _extract_pdf_text(data)
    elif is_docx:
        text = _extract_docx_text(data)
    else:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "Only PDF and DOCX assessment documents are supported")

    normalized = _normalize_text(text)
    if not normalized:
        raise HTTPException(
            status.HTTP_400_BAD_REQUEST,
            "No readable text was found in the document. Scanned PDFs need OCR before upload.",
        )

    return normalized[: settings.assessment_ai_max_input_chars]


async def generate_quiz_questions_from_text(
    *,
    source_text: str,
    question_count: int,
    options_per_question: int,
    course_title: str,
    section_title: str,
    item_title: str,
    provider: AssessmentAIProviderEnum,
    model: str | None,
) -> QuizAIGenerationResult:
    prompt = _build_document_generation_prompt(
        source_text=source_text,
        question_count=question_count,
        options_per_question=options_per_question,
        course_title=course_title,
        section_title=section_title,
        item_title=item_title,
    )
    return await _generate_quiz_questions(
        prompt=prompt,
        question_count=question_count,
        options_per_question=options_per_question,
        provider=provider,
        model=model,
    )


async def generate_quiz_questions_from_prompt(
    *,
    instructor_prompt: str,
    question_count: int,
    options_per_question: int,
    course_title: str,
    section_title: str,
    item_title: str,
    provider: AssessmentAIProviderEnum,
    model: str | None,
) -> QuizAIGenerationResult:
    prompt = _build_prompt_generation_prompt(
        instructor_prompt=instructor_prompt,
        question_count=question_count,
        options_per_question=options_per_question,
        course_title=course_title,
        section_title=section_title,
        item_title=item_title,
    )
    return await _generate_quiz_questions(
        prompt=prompt,
        question_count=question_count,
        options_per_question=options_per_question,
        provider=provider,
        model=model,
    )


async def _generate_quiz_questions(
    *,
    prompt: str,
    question_count: int,
    options_per_question: int,
    provider: AssessmentAIProviderEnum,
    model: str | None,
) -> QuizAIGenerationResult:
    if provider == AssessmentAIProviderEnum.GEMINI:
        selected_model = (model or settings.gemini_model).strip()
        response_text = await _call_gemini(prompt, question_count, options_per_question, selected_model)
    elif provider == AssessmentAIProviderEnum.OPENAI:
        selected_model = (model or settings.openai_model).strip()
        response_text = await _call_responses_api(
            provider_label="OpenAI",
            api_key=settings.openai_api_key,
            base_url=settings.openai_api_base_url,
            model=selected_model,
            prompt=prompt,
            question_count=question_count,
            options_per_question=options_per_question,
            include_strict=True,
        )
    elif provider == AssessmentAIProviderEnum.DEEPSEEK:
        selected_model = (model or settings.deepseek_model).strip()
        response_text = await _call_responses_api(
            provider_label="DeepSeek",
            api_key=settings.deepseek_api_key,
            base_url=settings.deepseek_api_base_url,
            model=selected_model,
            prompt=prompt,
            question_count=question_count,
            options_per_question=options_per_question,
            include_strict=False,
        )
    else:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "Unsupported assessment AI provider")

    validated = _validate_generation(response_text, question_count, options_per_question)
    return QuizAIGenerationResult(provider=provider, model=selected_model, questions=validated.questions)


async def _call_gemini(
    prompt: str, question_count: int, options_per_question: int, model: str
) -> str:
    if not settings.gemini_api_key:
        raise HTTPException(status.HTTP_500_INTERNAL_SERVER_ERROR, "GEMINI_API_KEY is not configured")

    payload = _build_gemini_payload(
        prompt=prompt,
        question_count=question_count,
        options_per_question=options_per_question,
    )

    normalized_model = _normalize_gemini_model_name(model)
    url = f"{settings.gemini_api_base_url.rstrip('/')}/{normalized_model}:generateContent"

    try:
        async with httpx.AsyncClient(timeout=settings.gemini_timeout_seconds) as client:
            response = await client.post(url, params={"key": settings.gemini_api_key}, json=payload)
            response.raise_for_status()
    except httpx.HTTPStatusError as exc:
        detail = _provider_error_message(exc.response)
        raise HTTPException(status.HTTP_502_BAD_GATEWAY, f"Gemini could not generate quiz questions: {detail}") from exc
    except httpx.HTTPError as exc:
        raise HTTPException(status.HTTP_502_BAD_GATEWAY, "Gemini request failed") from exc

    return _extract_gemini_text(response.json())


async def _call_responses_api(
    *,
    provider_label: str,
    api_key: str,
    base_url: str,
    model: str,
    prompt: str,
    question_count: int,
    options_per_question: int,
    include_strict: bool,
) -> str:
    if not api_key:
        env_name = "OPENAI_API_KEY" if provider_label == "OpenAI" else "DEEPSEEK_API_KEY"
        raise HTTPException(status.HTTP_500_INTERNAL_SERVER_ERROR, f"{env_name} is not configured")

    payload = _build_responses_payload(
        prompt=prompt,
        question_count=question_count,
        options_per_question=options_per_question,
        model=model,
        include_strict=include_strict,
    )
    url = f"{base_url.rstrip('/')}/responses"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

    try:
        async with httpx.AsyncClient(timeout=settings.gemini_timeout_seconds) as client:
            response = await client.post(url, headers=headers, json=payload)
            response.raise_for_status()
    except httpx.HTTPStatusError as exc:
        detail = _provider_error_message(exc.response)
        raise HTTPException(
            status.HTTP_502_BAD_GATEWAY,
            f"{provider_label} could not generate quiz questions: {detail}",
        ) from exc
    except httpx.HTTPError as exc:
        raise HTTPException(status.HTTP_502_BAD_GATEWAY, f"{provider_label} request failed") from exc

    return _extract_responses_text(response.json(), provider_label)


def _extract_pdf_text(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "Could not read the PDF assessment document") from exc


def _extract_docx_text(data: bytes) -> str:
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "Could not read the DOCX assessment document") from exc

    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _normalize_text(text: str) -> str:
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()


def _file_extension(file_name: str) -> str:
    lower_name = (file_name or "").lower()
    if lower_name.endswith(".pdf"):
        return ".pdf"
    if lower_name.endswith(".docx"):
        return ".docx"
    return ""


def _normalize_gemini_model_name(model: str) -> str:
    model = (model or "gemini-3.7-flash").strip()
    return model if model.startswith("models/") else f"models/{model}"


def _build_document_generation_prompt(
    *,
    source_text: str,
    question_count: int,
    options_per_question: int,
    course_title: str,
    section_title: str,
    item_title: str,
) -> str:
    return f"""
Create a high-quality multiple-choice quiz from the assessment document text.

Course: {course_title}
Module: {section_title}
Assessment item: {item_title}

Rules:
- Generate exactly {question_count} questions unless the document has too little material.
- Each question must test meaningful understanding, not wording trivia.
- Each question must have exactly {options_per_question} answer options.
- Most questions should have one correct answer. Use multiple correct answers only when the source clearly supports it.
- For a single-correct question, set allow_multiple_answers=false and exactly one option is_correct=true.
- For a multi-correct question, set allow_multiple_answers=true, multi_answer_mode="OR", and at least two options are is_correct=true.
- Do not invent facts that are not supported by the document.
- Keep wording clear for social work learners.

Assessment document text:
{source_text}
""".strip()


def _build_prompt_generation_prompt(
    *,
    instructor_prompt: str,
    question_count: int,
    options_per_question: int,
    course_title: str,
    section_title: str,
    item_title: str,
) -> str:
    return f"""
Create a high-quality multiple-choice quiz from the instructor's requested topics and learning goals.

Course: {course_title}
Module: {section_title}
Assessment item: {item_title}

Instructor request:
{instructor_prompt}

Rules:
- Generate exactly {question_count} questions unless the request has too little material.
- Each question must test meaningful understanding and practical application, not wording trivia.
- Each question must have exactly {options_per_question} answer options.
- Most questions should have one correct answer. Use multiple correct answers only when it improves the assessment.
- For a single-correct question, set allow_multiple_answers=false and exactly one option is_correct=true.
- For a multi-correct question, set allow_multiple_answers=true, multi_answer_mode="OR", and at least two options are is_correct=true.
- Stay within the instructor's requested areas. Where the request is broad, use standard social work education knowledge.
- Keep wording clear for social work learners.
""".strip()


def _build_gemini_payload(
    *,
    prompt: str,
    question_count: int,
    options_per_question: int,
) -> dict[str, Any]:
    return {
        "systemInstruction": {
            "parts": [
                {
                    "text": (
                        "You are an expert assessment designer. Return only valid JSON that matches the schema. "
                        "Do not include markdown, commentary, citations, or explanations."
                    )
                }
            ]
        },
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": 0.25,
            "responseMimeType": "application/json",
            "responseSchema": _quiz_response_schema(question_count, options_per_question),
        },
    }


def _build_responses_payload(
    *,
    prompt: str,
    question_count: int,
    options_per_question: int,
    model: str,
    include_strict: bool,
) -> dict[str, Any]:
    text_format = {
        "type": "json_schema",
        "name": "quiz_questions",
        "schema": _strict_quiz_response_schema(question_count, options_per_question),
    }
    if include_strict:
        text_format["strict"] = True

    return {
        "model": model,
        "instructions": (
            "You are an expert assessment designer. Return only valid JSON that matches the schema. "
            "Do not include markdown, commentary, citations, or explanations."
        ),
        "input": prompt,
        "text": {"format": text_format},
    }


def _quiz_response_schema(question_count: int, options_per_question: int) -> dict[str, Any]:
    return {
        "type": "object",
        "properties": {
            "questions": {
                "type": "array",
                "minItems": 1,
                "maxItems": question_count,
                "items": {
                    "type": "object",
                    "properties": {
                        "text": {"type": "string"},
                        "order_index": {"type": "integer"},
                        "allow_multiple_answers": {"type": "boolean"},
                        "multi_answer_mode": {"type": "string", "enum": ["AND", "OR"]},
                        "options": {
                            "type": "array",
                            "minItems": options_per_question,
                            "maxItems": options_per_question,
                            "items": {
                                "type": "object",
                                "properties": {
                                    "text": {"type": "string"},
                                    "is_correct": {"type": "boolean"},
                                    "order_index": {"type": "integer"},
                                },
                                "required": ["text", "is_correct", "order_index"],
                            },
                        },
                    },
                    "required": ["text", "order_index", "allow_multiple_answers", "options"],
                },
            }
        },
        "required": ["questions"],
    }


def _strict_quiz_response_schema(question_count: int, options_per_question: int) -> dict[str, Any]:
    return {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "questions": {
                "type": "array",
                "minItems": 1,
                "maxItems": question_count,
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "text": {"type": "string"},
                        "order_index": {"type": "integer"},
                        "allow_multiple_answers": {"type": "boolean"},
                        "multi_answer_mode": {"type": "string", "enum": ["AND", "OR"]},
                        "options": {
                            "type": "array",
                            "minItems": options_per_question,
                            "maxItems": options_per_question,
                            "items": {
                                "type": "object",
                                "additionalProperties": False,
                                "properties": {
                                    "text": {"type": "string"},
                                    "is_correct": {"type": "boolean"},
                                    "order_index": {"type": "integer"},
                                },
                                "required": ["text", "is_correct", "order_index"],
                            },
                        },
                    },
                    "required": [
                        "text",
                        "order_index",
                        "allow_multiple_answers",
                        "multi_answer_mode",
                        "options",
                    ],
                },
            }
        },
        "required": ["questions"],
    }


def _extract_gemini_text(payload: dict[str, Any]) -> str:
    candidates = payload.get("candidates") or []
    for candidate in candidates:
        content = candidate.get("content") or {}
        for part in content.get("parts") or []:
            text = part.get("text")
            if text:
                return text
    raise HTTPException(status.HTTP_502_BAD_GATEWAY, "Gemini returned no quiz content")


def _extract_responses_text(payload: dict[str, Any], provider_label: str) -> str:
    if payload.get("status") == "failed":
        error = payload.get("error") or {}
        detail = error.get("message") or "response failed"
        raise HTTPException(status.HTTP_502_BAD_GATEWAY, f"{provider_label} response failed: {detail}")

    for item in payload.get("output") or []:
        if item.get("type") != "message":
            continue
        for content in item.get("content") or []:
            if content.get("type") == "output_text" and content.get("text"):
                return content["text"]

    raise HTTPException(status.HTTP_502_BAD_GATEWAY, f"{provider_label} returned no quiz content")


def _validate_generation(response_text: str, question_count: int, options_per_question: int) -> RawQuizAIGenerationResult:
    try:
        raw = json.loads(_strip_json_fence(response_text))
    except json.JSONDecodeError as exc:
        raise HTTPException(status.HTTP_502_BAD_GATEWAY, "AI provider returned invalid JSON") from exc

    try:
        generated = RawQuizAIGenerationResult.model_validate(raw)
    except ValidationError as exc:
        raise HTTPException(status.HTTP_502_BAD_GATEWAY, "AI provider returned quiz data in an unexpected shape") from exc

    if not generated.questions:
        raise HTTPException(status.HTTP_502_BAD_GATEWAY, "AI provider did not generate any quiz questions")

    normalized_questions: list[QuizQuestionCreateDTO] = []
    for question_index, question in enumerate(generated.questions[:question_count]):
        options = question.options[:options_per_question]
        if len(options) < options_per_question:
            raise HTTPException(
                status.HTTP_502_BAD_GATEWAY,
                f"AI provider generated a question with fewer than {options_per_question} options",
            )

        correct_count = sum(1 for option in options if option.is_correct)
        if correct_count == 0:
            raise HTTPException(status.HTTP_502_BAD_GATEWAY, "AI provider generated a question with no correct answer")

        allow_multiple_answers = correct_count > 1
        normalized_options = [
            QuizOptionCreateDTO(
                text=option.text.strip(),
                is_correct=option.is_correct,
                order_index=option_index,
            )
            for option_index, option in enumerate(options)
        ]
        normalized_questions.append(
            QuizQuestionCreateDTO(
                text=question.text.strip(),
                order_index=question_index,
                allow_multiple_answers=allow_multiple_answers,
                multi_answer_mode=MultiAnswerModeEnum.OR if allow_multiple_answers else None,
                options=normalized_options,
            )
        )

    return RawQuizAIGenerationResult(questions=normalized_questions)


def _strip_json_fence(text: str) -> str:
    stripped = text.strip()
    if stripped.startswith("```"):
        stripped = re.sub(r"^```(?:json)?\s*", "", stripped, flags=re.IGNORECASE)
        stripped = re.sub(r"\s*```$", "", stripped)
    return stripped.strip()


def _provider_error_message(response: httpx.Response) -> str:
    try:
        payload = response.json()
    except ValueError:
        return response.text[:300] or response.reason_phrase
    error = payload.get("error") or {}
    return error.get("message") or response.reason_phrase
